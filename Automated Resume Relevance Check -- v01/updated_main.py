from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Optional
import io
import re
import logging
from datetime import datetime

# File parsing libraries
import PyPDF2
import docx

# LangChain and ChromaDB for vector operations
from langchain_ollama import OllamaEmbeddings
from langchain_core.documents import Document
from langchain_chroma import Chroma
import chromadb
from chromadb.config import Settings

# --- Configuration & Setup ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Intelligent Resume Ranker API",
    description="Implements a section-focused, two-layer system for highly accurate resume screening.",
    version="2.0.1"
)

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Pydantic Data Models for API Response ---
class RankedResume(BaseModel):
    rank: int
    filename: str
    semantic_score: float
    skill_match_count: int
    total_skills_required: int
    matched_skills: List[str]
    missing_skills: List[str]
    content_preview: str
    key_matching_chunk: str

class RankingResult(BaseModel):
    job_title: str
    total_resumes_processed: int
    top_candidates: List[RankedResume]
    processing_time: float
    summary: Dict

# --- Advanced Text Processing with Section Extraction ---
class TextProcessor:
    def clean_text(self, text: str) -> str:
        text = re.sub(r'[^\w\s\.,\-\+#@]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    def extract_text_from_file(self, content: bytes, filename: str) -> str:
        """Extract text from uploaded file with better error handling"""
        extension = filename.split('.')[-1].lower()
        text = ""
        
        try:
            if extension == 'pdf':
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                if reader.pages:
                    for page in reader.pages:
                        extracted = page.extract_text() or ""
                        text += extracted + "\n"
                else:
                    logger.warning(f"PDF {filename} has no readable pages (might be scanned)")
                    return ""
            
            elif extension == 'docx':
                doc = docx.Document(io.BytesIO(content))
                if doc.paragraphs:
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                else:
                    logger.warning(f"DOCX {filename} has no readable paragraphs")
                    return ""
            
            elif extension == 'txt':
                text = content.decode('utf-8', errors='ignore')
                if not text.strip():
                    logger.warning(f"TXT {filename} is empty")
                    return ""
            
            else:
                logger.warning(f"Unsupported file type: {filename}")
                return ""

            cleaned_text = self.clean_text(text)
            if not cleaned_text.strip():
                logger.warning(f"No extractable text found in {filename}")
                return ""
            
            return cleaned_text

        except Exception as e:
            logger.error(f"Failed to process {filename}: {e}")
            return ""

    def extract_sections(self, text: str) -> Dict[str, str]:
        """
        Extracts key sections from a resume using regex. This is a crucial step
        for targeted semantic search.
        """
        sections = {}
        
        # Define patterns for key sections. Case-insensitive, multiline.
        patterns = {
            'summary': r'(summary|objective|profile|about)\s*:?',
            'experience': r'(experience|employment history|work history|professional experience)\s*:?',
            'skills': r'(skills|technical skills|core competencies|technologies)\s*:?',
            'projects': r'(projects|personal projects|portfolio|key projects)\s*:?'
        }

        # Split text by lines for easier processing
        text_lower = text.lower()
        
        # Find start indices of all potential sections
        section_starts = {}
        for sec, pat in patterns.items():
            match = re.search(pat, text_lower, re.IGNORECASE)
            if match:
                section_starts[sec] = match.start()

        # Sort sections by their start position to handle them in order
        sorted_sections = sorted(section_starts.items(), key=lambda item: item[1])

        for i, (sec, start_index) in enumerate(sorted_sections):
            # The end index is the start of the next section, or the end of the text
            end_index = sorted_sections[i+1][1] if i + 1 < len(sorted_sections) else len(text)
            section_content = text[start_index:end_index].strip()
            
            # Clean up the header from the content
            header_match = re.match(patterns[sec], section_content, re.IGNORECASE | re.DOTALL)
            if header_match:
                section_content = section_content[header_match.end():].strip()
            
            if section_content:
                sections[sec] = section_content

        return sections

# --- Vector Store Manager ---
class VectorStoreManager:
    def __init__(self, embedding_model: OllamaEmbeddings):
        self.embedding_model = embedding_model
        self.chroma_client = chromadb.Client(Settings(anonymized_telemetry=False))
        self.collection_name = None
        self.vector_store = None

    def create_from_documents(self, docs: List[Document]):
        self.collection_name = f"resume_ranking_{datetime.now().timestamp()}".replace('.', '')[:63]
        
        try:
            self.chroma_client.delete_collection(name=self.collection_name)
        except Exception:
            pass

        self.vector_store = Chroma.from_documents(
            docs, 
            self.embedding_model, 
            client=self.chroma_client, 
            collection_name=self.collection_name
        )

    def similarity_search(self, query: str, k: int):
        if not self.vector_store:
            raise ValueError("Vector store has not been created.")
        return self.vector_store.similarity_search_with_score(query, k=k)

    def delete_collection(self):
        if self.collection_name:
            try:
                self.chroma_client.delete_collection(name=self.collection_name)
                self.collection_name = None
                self.vector_store = None
            except Exception as e:
                logger.warning(f"Could not delete collection {self.collection_name}: {e}")

# --- Core Resume Ranking Logic ---
class ResumeRanker:
    SEMANTIC_SEARCH_K = 20

    def __init__(self):
        try:
            self.embedding_model = OllamaEmbeddings(model="nomic-embed-text")
            logger.info("✅ Ollama Embeddings model initialized.")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Could not connect to Ollama: {e}")
        
        self.text_processor = TextProcessor()
        self.vector_store_manager = VectorStoreManager(self.embedding_model)

    def improved_skill_match(self, resume_text: str, skills: List[str]) -> tuple[List[str], List[str]]:
        """
        Improved skill matching with better regex boundaries
        Handles skills with symbols like C++, C# correctly
        """
        resume_text_lower = resume_text.lower()
        matched_skills = []
        
        for skill in skills:
            skill_lower = skill.strip().lower()
            if not skill_lower:
                continue
                
            # Use lookahead/lookbehind for better boundary detection
            # This handles C++, C#, .NET etc. properly
            pattern = r'(?<!\w)' + re.escape(skill_lower) + r'(?!\w)'
            
            if re.search(pattern, resume_text_lower, re.IGNORECASE):
                matched_skills.append(skill.strip())
        
        missing_skills = [skill for skill in skills if skill not in matched_skills]
        return matched_skills, missing_skills

    async def rank_resumes(self, job_description: str, required_skills: List[str], 
                          file_data_list: List[tuple], top_n: int = 10) -> RankingResult:
        """
        Made async and accepts pre-read file data
        """
        start_time = datetime.now()
        processed_resumes = {}
        section_documents = []
        
        try:
            # Pre-processing: Extract text and parse sections for all resumes
            for filename, content in file_data_list:
                if not content:
                    continue
                
                full_text = self.text_processor.extract_text_from_file(content, filename)
                if not full_text:
                    logger.warning(f"Could not extract text from {filename}, skipping.")
                    continue
                
                # Extract sections instead of generic chunks
                sections = self.text_processor.extract_sections(full_text)
                if not sections:
                    # Fallback: if no sections found, use full text as 'general' section
                    sections = {'general': full_text[:2000]}
                
                processed_resumes[filename] = {'full_text': full_text, 'sections': sections}
                
                # Create vectorizable documents ONLY from the extracted key sections
                for sec_name, sec_content in sections.items():
                    section_documents.append(Document(
                        page_content=sec_content,
                        metadata={'filename': filename, 'section': sec_name}
                    ))

            if not processed_resumes:
                raise HTTPException(status_code=400, detail="No valid resumes could be processed.")

            # Layer 1: Section-Focused Semantic Search
            logger.info(f"Starting Layer 1: Performing semantic search on {len(section_documents)} resume sections.")
            semantic_candidates = {}
            
            if section_documents:
                self.vector_store_manager.create_from_documents(section_documents)
                
                k = min(self.SEMANTIC_SEARCH_K, len(section_documents))
                similar_sections = self.vector_store_manager.similarity_search(job_description, k=k)
                
                # Aggregate scores: find the BEST matching section for each resume
                for doc, distance_score in similar_sections:
                    filename = doc.metadata['filename']
                    
                    if filename not in semantic_candidates or distance_score < semantic_candidates[filename]['raw_distance']:
                        semantic_candidates[filename] = {
                            'raw_distance': distance_score,
                            'chunk': f"[{doc.metadata.get('section', 'general').upper()}] {doc.page_content[:500]}..."
                        }

            logger.info(f"Layer 1 complete. Found {len(semantic_candidates)} unique candidates from section analysis.")

            # Layer 2: Skill-based Re-ranking on TOP semantic candidates
            re_ranking_list = []
            for filename, data in semantic_candidates.items():
                full_text = processed_resumes[filename]['full_text']
                
                matched_skills, missing_skills = self.improved_skill_match(full_text, required_skills)
                
                re_ranking_list.append({
                    'filename': filename,
                    'raw_distance': data['raw_distance'],
                    'skill_match_count': len(matched_skills),
                    'matched_skills': matched_skills,
                    'missing_skills': missing_skills,
                    'key_matching_chunk': data['chunk']
                })

            # Sort by skill count (desc), then distance (asc - lower is better)
            re_ranking_list.sort(key=lambda x: (-x['skill_match_count'], x['raw_distance']))

            # Prepare response with normalized scores for frontend compatibility
            top_candidates = []
            for i, c in enumerate(re_ranking_list[:top_n]):
                # Convert distance to normalized score for frontend
                normalized_score = max(0.0, min(1.0, 1.0 / (1.0 + c['raw_distance'])))
                
                top_candidates.append(RankedResume(
                    rank=i + 1,
                    filename=c['filename'],
                    semantic_score=normalized_score,
                    skill_match_count=c['skill_match_count'],
                    total_skills_required=len(required_skills),
                    matched_skills=c['matched_skills'],
                    missing_skills=c['missing_skills'],
                    content_preview=processed_resumes[c['filename']]['full_text'][:400] + "...",
                    key_matching_chunk=c.get('key_matching_chunk', "N/A")
                ))

            processing_time = (datetime.now() - start_time).total_seconds()
            summary = {
                "Resumes Processed": len(processed_resumes),
                "Semantic Candidates Found": len(semantic_candidates),
                "Final Recommendations": len(top_candidates)
            }

            return RankingResult(
                job_title="",
                total_resumes_processed=len(processed_resumes),
                top_candidates=top_candidates,
                processing_time=processing_time,
                summary=summary
            )

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"An unexpected error occurred during ranking: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"An internal error occurred: {str(e)}")
        finally:
            self.vector_store_manager.delete_collection()

# --- Global Singleton Instance ---
ranker = ResumeRanker()

# --- API Endpoints ---
@app.get("/")
def root():
    return {"message": "Resume Ranking API is ready.", "version": "2.0.1"}

@app.post("/rank-resumes", response_model=RankingResult)
async def rank_resumes_endpoint(
    job_title: Optional[str] = Form(None),
    job_description: str = Form(...),
    required_skills: str = Form(...),
    top_n: int = Form(10),
    files: List[UploadFile] = File(...)
):
    # Parse and clean skills
    cleaned_skills_str = re.sub(r'[()]', ' ', required_skills)
    cleaned_skills_str = cleaned_skills_str.replace(' or ', ',').replace(' and ', ',')
    skill_list = [skill.strip() for skill in cleaned_skills_str.split(',') if skill.strip()]

    if not job_description or not skill_list:
        raise HTTPException(status_code=422, detail="Job Description and Required Skills cannot be empty.")

    # Async file reading
    file_data_list = []
    for file in files:
        try:
            content = await file.read()
            if content:
                file_data_list.append((file.filename, content))
        except Exception as e:
            logger.error(f"Error reading file {file.filename}: {e}")
            continue

    if not file_data_list:
        raise HTTPException(status_code=400, detail="No valid files could be read.")

    try:
        result = await ranker.rank_resumes(job_description, skill_list, file_data_list, top_n)
        result.job_title = job_title if job_title else "AI Analysis Result"
        return result
    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"An unexpected error occurred in endpoint: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An internal error occurred: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)